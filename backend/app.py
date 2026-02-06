from forecast import forecast_revenue
from bookkeeping import auto_bookkeeping
from benchmark import compare_with_benchmark
from fastapi.responses import StreamingResponse
from pdf_report import generate_pdf_report
from recommendation import recommend_products
from tax_rules import tax_compliance_checks
from tax import check_tax_compliance, tax_compliance_checks
from fastapi import FastAPI, UploadFile, Form, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
import io

from analysis import analyze_financials
from database import SessionLocal, engine
from models import Base, AnalysisResult
from llm import generate_financial_report, translate_to_hindi

from docx import Document
import pdfplumber

# create tables
Base.metadata.create_all(bind=engine)

app = FastAPI(title="SME Financial Health Assessment API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# DB dependency
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


@app.get("/")
def root():
    return {"status": "SME Financial Health API running"}

def parse_text_lines_to_df(lines):
    rows = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # try comma separated
        if "," in line:
            parts = [p.strip() for p in line.split(",")]
        else:
            # fallback: whitespace separated
            parts = line.split()

        if len(parts) >= 2:
            rows.append([parts[0], parts[1]])

    if not rows:
        return None

    # remove header if present
    if rows[0][0].lower() == "amount":
        rows = rows[1:]

    df = pd.DataFrame(rows, columns=["amount", "type"])
    df["amount"] = pd.to_numeric(df["amount"], errors="coerce")

    return df

# -------- helpers for DOCX and PDF --------

def read_docx_to_dataframe(file_bytes: bytes) -> pd.DataFrame:
    doc = Document(io.BytesIO(file_bytes))

    rows = []

    # -------- 1. Try tables first --------
    for table in doc.tables:
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                rows.append([cells[0], cells[1]])

    if rows:
        df = pd.DataFrame(rows, columns=["amount", "type"])
        df["amount"] = pd.to_numeric(df["amount"], errors="coerce")
        return df

    # -------- 2. Fallback: normal text lines --------
    text_rows = []

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue

        # expect: amount,type
        if "," in line:
            parts = line.split(",")
            if len(parts) >= 2:
                text_rows.append([parts[0].strip(), parts[1].strip()])

    if not text_rows:
        raise ValueError(
            "No table or valid text lines found in Word document. "
            "Use a table or lines like: 50000,Revenue"
        )

    # remove header if present
    if text_rows[0][0].lower() == "amount":
        text_rows = text_rows[1:]

    df = pd.DataFrame(text_rows, columns=["amount", "type"])
    df["amount"] = pd.to_numeric(df["amount"], errors="coerce")

    return df



def read_pdf_to_dataframe(file_bytes: bytes) -> pd.DataFrame:
    rows = []

    all_text_lines = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:

            # -------- try tables first --------
            tables = page.extract_tables()
            for table in tables:
                for row in table[1:]:
                    if row and len(row) >= 2:
                        rows.append([row[0], row[1]])

            # -------- collect text for fallback --------
            text = page.extract_text()
            if text:
                all_text_lines.extend(text.splitlines())

    # ---- if table data found
    if rows:
        df = pd.DataFrame(rows, columns=["amount", "type"])
        df["amount"] = pd.to_numeric(df["amount"], errors="coerce")
        return df

    # ---- fallback to normal text
    df = parse_text_lines_to_df(all_text_lines)

    if df is None:
        raise ValueError(
            "No table or valid text lines found in PDF. "
            "Use table or lines like: 50000,Revenue"
        )

    return df



# -------------------------------------------

@app.post("/analyze")
async def analyze_file(
    file: UploadFile,
    business_name: str = Form(...),
    industry: str = Form(...),
    db=Depends(get_db)
):

    try:
        filename = file.filename.lower()
        content = await file.read()

        # -----------------------------
        # Read file
        # -----------------------------
        if filename.endswith(".csv"):
            df = pd.read_csv(io.BytesIO(content))

        elif filename.endswith(".xlsx") or filename.endswith(".xls"):
            df = pd.read_excel(io.BytesIO(content))

        elif filename.endswith(".docx"):
            df = read_docx_to_dataframe(content)

        elif filename.endswith(".pdf"):
            df = read_pdf_to_dataframe(content)

        else:
            raise HTTPException(
                status_code=400,
                detail="Only CSV, Excel, Word (.docx) and PDF files are supported."
            )

        # -----------------------------
        # Basic validation
        # -----------------------------
        required_cols = {"amount", "type"}
        df.columns = [c.lower() for c in df.columns]
        

        if not required_cols.issubset(set(df.columns)):
            raise HTTPException(
                status_code=400,
                detail="Input file must contain 'amount' and 'type' columns."
            )
        
        df["amount"] = pd.to_numeric(df["amount"], errors="coerce")
        df["type"] = df["type"].astype(str).str.strip().str.lower()
        full_df = df.copy()

        # -----------------------------
        # Analytics
        # -----------------------------
        analysis_df = df[["amount", "type"]].dropna()
        metrics, risks, score = analyze_financials(analysis_df)
        forecast = forecast_revenue(full_df)
        bookkeeping = auto_bookkeeping(full_df)
        benchmark = compare_with_benchmark(metrics, industry)
        products = recommend_products(metrics, score)
        tax_warnings = tax_compliance_checks(full_df, metrics)
        tax_compliance = check_tax_compliance(full_df)

        tax_compliance["warnings"] = tax_warnings



        # -----------------------------
        # AI Report
        # -----------------------------
        try:
            english_report = generate_financial_report(
                business_name=business_name,
                industry=industry,
                metrics=metrics,
                risks=risks,
                credit_score=score
            )

            hindi_report = translate_to_hindi(english_report)

        except Exception as e:
            english_report = (
                "AI report is temporarily unavailable due to API quota limits. "
                "Financial metrics and risk analysis were generated successfully."
            )
            hindi_report = (
                "एआई रिपोर्ट अस्थायी रूप से उपलब्ध नहीं है। "
                "वित्तीय मेट्रिक्स और जोखिम विश्लेषण सफलतापूर्वक तैयार किए गए हैं।"
            )

        # -----------------------------
        # Store in DB
        # -----------------------------
        row = AnalysisResult(
            business_name=business_name,
            industry=industry,
            metrics=metrics,
            risks=risks,
            credit_score=score,
            ai_report=english_report
        )

        db.add(row)
        db.commit()
        db.refresh(row)

        return {
            "id": row.id,
            "business_name": business_name,
            "industry": industry,
            "metrics": metrics,
            "risks": risks,
            "credit_score": score,
            "forecast": forecast,
            "bookkeeping": bookkeeping,
            "benchmark": benchmark,
            "recommended_products": products,
            "tax_compliance": tax_compliance,
            "report_en": english_report,
            "report_hi": hindi_report
            
        }

    except HTTPException:
        raise

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )
@app.get("/report/pdf/{analysis_id}")
def download_pdf(analysis_id: int, db=Depends(get_db)):

    row = db.query(AnalysisResult).filter(AnalysisResult.id == analysis_id).first()

    if not row:
        raise HTTPException(status_code=404, detail="Not found")

    data = {
        "business_name": row.business_name,
        "industry": row.industry,
        "metrics": row.metrics,
        "credit_score": row.credit_score,
        "report_en": row.ai_report
    }

    pdf = generate_pdf_report(data)

    return StreamingResponse(
        iter([pdf]),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=report_{analysis_id}.pdf"
        }
    )
